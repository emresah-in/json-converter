import json
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
import sys
from io import BytesIO
import logging

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def decode_json_data(data):
    """Handle various JSON input formats."""
    try:
        if isinstance(data, str):
            # If it's a string, try to parse it
            try:
                # First try direct JSON parsing
                return json.loads(data)
            except json.JSONDecodeError:
                # If that fails, try handling escaped JSON string
                if data.startswith('"') and data.endswith('"'):
                    # Remove outer quotes and unescape
                    inner_json = data[1:-1].encode().decode('unicode_escape')
                    return json.loads(inner_json)
                raise
        return data
    except Exception as e:
        logger.error(f"Error decoding JSON data: {str(e)}")
        raise

def read_json_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return decode_json_data(file.read())
    except Exception as e:
        logger.error(f"Error reading JSON file: {str(e)}")
        return None

def create_word_document(json_data, output):
    try:
        # Ensure we have proper JSON data
        json_data = decode_json_data(json_data)
        logger.debug(f"Processing JSON data: {str(json_data)[:200]}...")
        
        doc = Document()
        
        # Add document styling
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        
        # Add title if ID exists
        if 'id' in json_data:
            doc.add_heading(f"BIM Execution Plan (BEP)", level=1)
            doc.add_paragraph(f"Document ID: {json_data['id']}")
        
        # Add creation date if exists
        if 'createdAt' in json_data:
            doc.add_paragraph(f"Created: {json_data['createdAt']}")
            doc.add_paragraph()  # Add spacing
        
        # Process all other fields
        for key, value in json_data.items():
            # Skip id and createdAt as they're already handled
            if key in ['id', 'createdAt']:
                continue
                
            # Skip null values
            if value is None or value == "null":
                continue
            
            # Convert key from camelCase to Title Case for heading
            heading_text = ''.join(' ' + char if char.isupper() else char for char in key).strip()
            heading_text = heading_text.title()
            
            # Add section heading
            doc.add_heading(heading_text, level=2)
            
            # Add content with proper paragraph formatting
            paragraph = doc.add_paragraph(str(value))
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            paragraph.paragraph_format.space_after = Pt(12)  # Add spacing after paragraph
        
        logger.debug("Saving Word document...")
        # Save to file or BytesIO
        if isinstance(output, str):
            doc.save(output)
        else:
            doc.save(output)
        logger.debug("Word document saved successfully")
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise

def create_pdf_document(json_data, output):
    try:
        # Ensure we have proper JSON data
        json_data = decode_json_data(json_data)
        logger.debug(f"Processing JSON data: {str(json_data)[:200]}...")
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Set default font
        pdf.set_font("Times", "", 12)
        
        # Add title
        pdf.set_font("Times", "B", 16)
        pdf.cell(0, 10, "BIM Execution Plan (BEP)", ln=True, align='C')
        pdf.ln(5)
        
        # Add ID if exists
        if 'id' in json_data:
            pdf.set_font("Times", "", 12)
            pdf.cell(0, 10, f"Document ID: {json_data['id']}", ln=True)
        
        # Add creation date if exists
        if 'createdAt' in json_data:
            pdf.set_font("Times", "", 12)
            pdf.cell(0, 10, f"Created: {json_data['createdAt']}", ln=True)
            pdf.ln(5)
        
        # Process all other fields
        for key, value in json_data.items():
            # Skip id and createdAt as they're already handled
            if key in ['id', 'createdAt']:
                continue
                
            # Skip null values
            if value is None or value == "null":
                continue
            
            # Convert key from camelCase to Title Case for heading
            heading_text = ''.join(' ' + char if char.isupper() else char for char in key).strip()
            heading_text = heading_text.title()
            
            # Add section heading
            pdf.set_font("Times", "B", 14)
            pdf.cell(0, 10, heading_text, ln=True)
            
            # Add content
            pdf.set_font("Times", "", 12)
            pdf.multi_cell(0, 10, str(value))
            pdf.ln(5)  # Add spacing between sections
        
        logger.debug("Saving PDF document...")
        # Save to file or BytesIO
        if isinstance(output, str):
            pdf.output(output)
        else:
            pdf.output(output)
        logger.debug("PDF document saved successfully")
        
    except Exception as e:
        logger.error(f"Error creating PDF document: {str(e)}")
        raise

def convert_json(json_file_path, output_format="both"):
    # Read JSON file
    json_data = read_json_file(json_file_path)
    if not json_data:
        return False
    
    # Generate output file names
    file_base = json_file_path.rsplit('.', 1)[0]
    success = True
    
    # Convert based on requested format
    if output_format.lower() in ["both", "word"]:
        word_output = f"{file_base}.docx"
        try:
            create_word_document(json_data, word_output)
        except Exception as e:
            logger.error(f"Error creating Word document: {str(e)}")
            success = False
    
    if output_format.lower() in ["both", "pdf"]:
        pdf_output = f"{file_base}.pdf"
        try:
            create_pdf_document(json_data, pdf_output)
        except Exception as e:
            logger.error(f"Error creating PDF document: {str(e)}")
            success = False
    
    return success

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python json_converter.py <json_file_path> [output_format]")
        print("output_format can be: 'word', 'pdf', or 'both' (default)")
        sys.exit(1)
    
    json_file_path = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else "both"
    
    convert_json(json_file_path, output_format)
